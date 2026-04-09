from __future__ import annotations

from datetime import datetime, timezone
import csv
import io
import json
import logging
from pathlib import Path
from typing import Any
from uuid import uuid4
import zipfile

from pyproj import Transformer
from shapely.geometry import MultiPolygon, Polygon, shape
from shapely.ops import transform

from integracoes.arquivos_projeto import validar_lote_uploads, validar_upload_formulario
from integracoes.projeto_clientes import listar_participantes_area, salvar_participantes_area
from integracoes.referencia_cliente import importar_vertices_por_formato, resumir_vertices

try:
    from docx import Document
except Exception:  # pragma: no cover - fallback defensivo
    Document = None


BASE_DIR = Path(__file__).resolve().parent.parent
DATA_DIR = BASE_DIR / "data"
UPLOADS_DIR = DATA_DIR / "formulario_uploads"
TEMPLATES_DIR = BASE_DIR / "static" / "templates"
TEMPLATE_CARTA_CONFRONTACAO = TEMPLATES_DIR / "carta_confrontacao.docx"
logger = logging.getLogger("geoadmin.areas_projeto")

STATUS_OPERACIONAL_VALIDOS = {
    "aguardando_cliente",
    "cliente_vinculado",
    "croqui_recebido",
    "geometria_final",
    "peca_pronta",
}
STATUS_DOCUMENTAL_VALIDOS = {
    "pendente",
    "formulario_ok",
    "confrontantes_ok",
    "documentacao_ok",
    "peca_pronta",
}



def _get_supabase():
    from main import get_supabase
    return get_supabase()


def _agora_iso() -> str:
    return datetime.now(timezone.utc).isoformat()


def _normalizar_status_operacional(valor: str | None, *, possui_participante: bool = False, status_geometria: str | None = None) -> str:
    chave = (valor or '').strip().lower()
    if chave in STATUS_OPERACIONAL_VALIDOS:
        return chave
    if status_geometria == 'geometria_final':
        return 'geometria_final'
    if status_geometria == 'apenas_esboco':
        return 'croqui_recebido'
    if possui_participante:
        return 'cliente_vinculado'
    return 'aguardando_cliente'


def _normalizar_status_documental(valor: str | None, *, possui_participante: bool = False) -> str:
    chave = (valor or '').strip().lower()
    if chave in STATUS_DOCUMENTAL_VALIDOS:
        return chave
    return 'formulario_ok' if possui_participante else 'pendente'


def _identificacao_lote(area: dict[str, Any]) -> str:
    partes = []
    if area.get('quadra'):
        partes.append(f"Qd. {area['quadra']}")
    if area.get('codigo_lote'):
        partes.append(f"Lt. {area['codigo_lote']}")
    if area.get('setor'):
        partes.append(str(area['setor']))
    return ' · '.join(partes)


def _falhar_supabase(exc: Exception, operacao: str) -> None:
    raise RuntimeError(f"Falha ao {operacao} em areas_projeto no Supabase: {exc}") from exc


def _data_extenso_ptbr(data: datetime | None = None) -> str:
    meses = [
        "janeiro",
        "fevereiro",
        "março",
        "abril",
        "maio",
        "junho",
        "julho",
        "agosto",
        "setembro",
        "outubro",
        "novembro",
        "dezembro",
    ]
    valor = data or datetime.now()
    return f"{valor.day} de {meses[valor.month - 1]} de {valor.year}"


def _vertices_validos(vertices: list[dict[str, Any]] | None) -> list[dict[str, float]]:
    lista = vertices or []
    if len(lista) < 3:
        return []
    normalizados: list[dict[str, float]] = []
    for item in lista:
        normalizados.append({"lon": float(item["lon"]), "lat": float(item["lat"])})
    return normalizados


def _polygon_from_vertices(vertices: list[dict[str, Any]]) -> Polygon | None:
    coords = [(float(item["lon"]), float(item["lat"])) for item in vertices]
    if len(coords) < 3:
        return None
    polygon = Polygon(coords)
    if not polygon.is_valid:
        polygon = polygon.buffer(0)
    if polygon.is_empty:
        return None
    return polygon


def _utm_epsg(lat: float, lon: float) -> int:
    fuso = int((lon + 180) // 6) + 1
    return 32700 + fuso if lat < 0 else 32600 + fuso


def _transformar_para_metros(polygon: Polygon) -> Polygon:
    centroide = polygon.centroid
    epsg = _utm_epsg(float(centroide.y), float(centroide.x))
    transformer = Transformer.from_crs("EPSG:4326", f"EPSG:{epsg}", always_xy=True)
    return transform(transformer.transform, polygon)


def _resumo_vertices(vertices: list[dict[str, Any]] | None) -> dict[str, Any] | None:
    vertices_ok = _vertices_validos(vertices)
    if not vertices_ok:
        return None
    return resumir_vertices(vertices_ok)


def _geometria_preferencial(area: dict[str, Any]) -> tuple[str, list[dict[str, Any]]]:
    geometria_final = _vertices_validos(area.get("geometria_final"))
    if geometria_final:
        return "final", geometria_final
    esboco = _vertices_validos(area.get("geometria_esboco"))
    return "esboco", esboco


def _status_geometria(area: dict[str, Any]) -> str:
    if _vertices_validos(area.get("geometria_final")):
        return "geometria_final"
    if _vertices_validos(area.get("geometria_esboco")):
        return "apenas_esboco"
    return "sem_geometria"


def _slug_lote(valor: str | None) -> str:
    return (valor or '').strip().lower()


def _chave_lote(area: dict[str, Any]) -> str | None:
    codigo = _slug_lote(area.get('codigo_lote'))
    quadra = _slug_lote(area.get('quadra'))
    setor = _slug_lote(area.get('setor'))
    if not any((codigo, quadra, setor)):
        return None
    return f'{quadra}::{codigo}::{setor}'


def _bool_input(valor: Any, default: bool = False) -> bool:
    if valor is None:
        return default
    if isinstance(valor, bool):
        return valor
    texto = str(valor).strip().lower()
    if texto in {'1', 'true', 'sim', 's', 'yes', 'y'}:
        return True
    if texto in {'0', 'false', 'nao', 'não', 'n', 'no'}:
        return False
    return default


def _vertices_geojson_geometry(geometry: dict[str, Any] | None) -> list[dict[str, float]]:
    if not geometry:
        return []
    geom = shape(geometry)
    candidatos: list[Polygon] = []
    if isinstance(geom, Polygon):
        candidatos = [geom]
    elif isinstance(geom, MultiPolygon):
        candidatos = [item for item in geom.geoms if not item.is_empty]
    if not candidatos:
        return []
    maior = max(candidatos, key=lambda item: item.area)
    coords = list(maior.exterior.coords)
    if len(coords) > 1 and coords[0] == coords[-1]:
        coords = coords[:-1]
    return [{'lon': float(lon), 'lat': float(lat)} for lon, lat in coords]


def _participantes_importacao(bruto: dict[str, Any]) -> list[dict[str, Any]]:
    participantes = bruto.get('participantes_area') or bruto.get('participantes') or []
    if participantes:
        return participantes

    nome = bruto.get('participante_nome') or bruto.get('cliente_nome') or bruto.get('proprietario_nome')
    cpf = bruto.get('participante_cpf') or bruto.get('cliente_cpf')
    telefone = bruto.get('participante_telefone') or bruto.get('cliente_telefone')
    cliente_id = bruto.get('participante_cliente_id') or bruto.get('cliente_id')
    if not any((nome, cpf, telefone, cliente_id)):
        return []
    return [{
        'cliente_id': cliente_id,
        'nome': nome,
        'cpf': cpf,
        'telefone': telefone,
        'papel': bruto.get('papel') or 'principal',
        'principal': _bool_input(bruto.get('principal'), True),
        'recebe_magic_link': _bool_input(bruto.get('recebe_magic_link'), True),
        'ordem': 0,
    }]


def _normalizar_lote_importado(bruto: dict[str, Any], indice: int) -> dict[str, Any]:
    vertices_final = _vertices_validos(bruto.get('geometria_final') or bruto.get('vertices') or bruto.get('vertices_json'))
    vertices_esboco = _vertices_validos(bruto.get('geometria_esboco'))
    nome = (bruto.get('nome') or '').strip() or None
    codigo_lote = (bruto.get('codigo_lote') or bruto.get('lote') or '').strip() or None
    quadra = (bruto.get('quadra') or '').strip() or None
    setor = (bruto.get('setor') or '').strip() or None
    identificador = ' / '.join([item for item in (quadra and f'Qd. {quadra}', codigo_lote and f'Lt. {codigo_lote}', setor) if item])
    if not nome:
        nome = identificador or f'Lote {indice + 1}'

    participantes_area = _participantes_importacao(bruto)
    cliente_id = bruto.get('cliente_id') or next((item.get('cliente_id') for item in participantes_area if item.get('cliente_id')), None)
    proprietario = bruto.get('proprietario_nome') or next((item.get('nome') for item in participantes_area if item.get('nome')), None)
    return {
        'area_id': bruto.get('area_id') or bruto.get('id'),
        'nome': nome,
        'cliente_id': cliente_id,
        'proprietario_nome': proprietario,
        'municipio': bruto.get('municipio'),
        'estado': bruto.get('estado'),
        'comarca': bruto.get('comarca'),
        'matricula': bruto.get('matricula'),
        'ccir': bruto.get('ccir'),
        'car': bruto.get('car'),
        'observacoes': bruto.get('observacoes'),
        'codigo_lote': codigo_lote,
        'quadra': quadra,
        'setor': setor,
        'status_operacional': bruto.get('status_operacional'),
        'status_documental': bruto.get('status_documental'),
        'origem_tipo': bruto.get('origem_tipo') or 'importacao',
        'geometria_esboco': vertices_esboco,
        'geometria_final': vertices_final,
        'participantes_area': participantes_area,
    }


def parse_lotes_geojson(conteudo: str | dict[str, Any]) -> list[dict[str, Any]]:
    payload = json.loads(conteudo) if isinstance(conteudo, str) else conteudo
    if not isinstance(payload, dict):
        raise ValueError('GeoJSON invalido para importacao de lotes.')

    tipo = payload.get('type')
    if tipo == 'FeatureCollection':
        features = payload.get('features', [])
    elif tipo == 'Feature':
        features = [payload]
    else:
        features = [{'type': 'Feature', 'properties': {}, 'geometry': payload}]

    lotes: list[dict[str, Any]] = []
    for indice, feature in enumerate(features):
        geometry = feature.get('geometry')
        vertices = _vertices_geojson_geometry(geometry)
        if not vertices:
            continue
        props = feature.get('properties') or {}
        lotes.append(_normalizar_lote_importado({
            **props,
            'geometria_final': vertices,
        }, indice))

    if not lotes:
        raise ValueError('Nenhum lote valido encontrado no GeoJSON.')
    return lotes


def parse_lotes_csv(conteudo: str) -> list[dict[str, Any]]:
    amostra = '\n'.join(conteudo.splitlines()[:5])
    try:
        dialect = csv.Sniffer().sniff(amostra or 'codigo_lote,nome')
    except Exception:
        dialect = csv.excel

    reader = csv.DictReader(io.StringIO(conteudo), dialect=dialect)
    if not reader.fieldnames:
        raise ValueError('CSV sem cabeçalho para importacao de lotes.')

    lotes: list[dict[str, Any]] = []
    for indice, row in enumerate(reader):
        bruto = {
            chave.strip(): (valor.strip() if isinstance(valor, str) else valor)
            for chave, valor in row.items() if chave
        }
        vertices_json = bruto.get('vertices_json') or bruto.get('geometria_json') or bruto.get('geometria_final')
        if isinstance(vertices_json, str) and vertices_json.strip():
            try:
                bruto['geometria_final'] = json.loads(vertices_json)
            except Exception:
                bruto['geometria_final'] = []
        lotes.append(_normalizar_lote_importado(bruto, indice))

    if not lotes:
        raise ValueError('Nenhuma linha valida encontrada no CSV de lotes.')
    return lotes


def importar_lotes_por_formato(formato: str, conteudo: str | bytes) -> dict[str, Any]:
    chave = (formato or '').strip().lower()
    if chave in {'geojson', 'json'}:
        lotes = parse_lotes_geojson(conteudo.decode('utf-8') if isinstance(conteudo, bytes) else conteudo)
        return {'lotes': lotes, 'parcial': False, 'mensagem': None}
    if chave == 'csv':
        lotes = parse_lotes_csv(conteudo.decode('utf-8') if isinstance(conteudo, bytes) else conteudo)
        return {'lotes': lotes, 'parcial': False, 'mensagem': None}
    if chave in {'kml', 'zip', 'shpzip', 'txt'}:
        vertices = importar_vertices_por_formato(chave, conteudo)
        lote = _normalizar_lote_importado({'geometria_final': vertices, 'origem_tipo': 'importacao'}, 0)
        return {
            'lotes': [lote],
            'parcial': True,
            'mensagem': 'O parser atual reaproveitou apenas a geometria principal do arquivo. Para importacao massiva de lotes, prefira GeoJSON ou CSV estruturado.',
        }
    raise ValueError(f'Formato de importacao de lotes nao suportado: {formato}')


def importar_areas_projeto_em_lote(
    *,
    projeto_id: str,
    lotes: list[dict[str, Any]],
    atualizar_existentes: bool = True,
    sb=None,
) -> dict[str, Any]:
    cliente = sb or _get_supabase()
    existentes = listar_areas_projeto(projeto_id, sb=cliente)
    por_id = {str(area.get('id')): area for area in existentes if area.get('id')}
    por_chave = {chave: area for area in existentes if (chave := _chave_lote(area))}

    criadas = 0
    atualizadas = 0
    ignoradas = 0
    areas_salvas: list[dict[str, Any]] = []

    for indice, bruto in enumerate(lotes):
        lote = _normalizar_lote_importado(bruto, indice)
        existente = None
        if lote.get('area_id') and str(lote['area_id']) in por_id:
            existente = por_id[str(lote['area_id'])]
        elif (chave := _chave_lote(lote)) and chave in por_chave:
            existente = por_chave[chave]

        if existente and not atualizar_existentes:
            ignoradas += 1
            continue

        area_salva = salvar_area_projeto(
            projeto_id=projeto_id,
            cliente_id=lote.get('cliente_id'),
            nome=lote.get('nome') or 'Area sem nome',
            proprietario_nome=lote.get('proprietario_nome'),
            municipio=lote.get('municipio'),
            estado=lote.get('estado'),
            comarca=lote.get('comarca'),
            matricula=lote.get('matricula'),
            ccir=lote.get('ccir'),
            car=lote.get('car'),
            observacoes=lote.get('observacoes'),
            codigo_lote=lote.get('codigo_lote'),
            quadra=lote.get('quadra'),
            setor=lote.get('setor'),
            status_operacional=lote.get('status_operacional'),
            status_documental=lote.get('status_documental'),
            origem_tipo=lote.get('origem_tipo') or 'importacao',
            geometria_esboco=lote.get('geometria_esboco') or [],
            geometria_final=lote.get('geometria_final') or [],
            participantes_area=lote.get('participantes_area') or [],
            area_id=(existente or {}).get('id') or lote.get('area_id'),
            sb=cliente,
        )
        if existente:
            atualizadas += 1
        else:
            criadas += 1
        areas_salvas.append(area_salva)

    painel = montar_painel_lotes(areas_salvas)
    return {
        'total_recebido': len(lotes),
        'criadas': criadas,
        'atualizadas': atualizadas,
        'ignoradas': ignoradas,
        'areas': areas_salvas,
        'painel_lotes': painel,
    }


def montar_painel_lotes(areas: list[dict[str, Any]]) -> list[dict[str, Any]]:
    painel: list[dict[str, Any]] = []
    for area in areas:
        participantes = area.get('participantes_area') or []
        principal = next((item for item in participantes if item.get('principal')), None) or (participantes[0] if participantes else None)
        formulario_ok = any(bool(item.get('formulario_ok')) for item in participantes)
        formulario_em = next((item.get('formulario_em') for item in participantes if item.get('formulario_em')), None)
        painel.append({
            'area_id': area.get('id'),
            'nome': area.get('nome'),
            'codigo_lote': area.get('codigo_lote'),
            'quadra': area.get('quadra'),
            'setor': area.get('setor'),
            'identificacao_lote': area.get('identificacao_lote') or _identificacao_lote(area) or area.get('nome'),
            'status_operacional': area.get('status_operacional') or 'aguardando_cliente',
            'status_documental': area.get('status_documental') or 'pendente',
            'status_geometria': area.get('status_geometria') or _status_geometria(area),
            'participantes_total': len(participantes),
            'principal_nome': (principal or {}).get('nome') or area.get('proprietario_nome'),
            'principal_cpf': (principal or {}).get('cpf'),
            'recebe_magic_link_total': sum(1 for item in participantes if item.get('recebe_magic_link')),
            'formulario_ok': formulario_ok,
            'formulario_em': formulario_em,
        })
    painel.sort(key=lambda item: ((item.get('quadra') or ''), (item.get('codigo_lote') or ''), (item.get('nome') or '')))
    return painel


def _row_para_area(raw: dict[str, Any] | None) -> dict[str, Any] | None:
    if not raw:
        return None
    return {
        "id": raw.get("id"),
        "projeto_id": raw.get("projeto_id"),
        "cliente_id": raw.get("cliente_id"),
        "nome": raw.get("nome"),
        "proprietario_nome": raw.get("proprietario_nome"),
        "municipio": raw.get("municipio"),
        "estado": raw.get("estado"),
        "comarca": raw.get("comarca"),
        "matricula": raw.get("matricula"),
        "ccir": raw.get("ccir"),
        "car": raw.get("car"),
        "observacoes": raw.get("observacoes"),
        "codigo_lote": raw.get("codigo_lote"),
        "quadra": raw.get("quadra"),
        "setor": raw.get("setor"),
        "status_operacional": raw.get("status_operacional"),
        "status_documental": raw.get("status_documental"),
        "origem_tipo": raw.get("origem_tipo") or "manual",
        "geometria_esboco": raw.get("geometria_esboco") or [],
        "geometria_final": raw.get("geometria_final") or [],
        "resumo_esboco": raw.get("resumo_esboco"),
        "resumo_final": raw.get("resumo_final"),
        "anexos": raw.get("anexos") or [],
        "criado_em": raw.get("criado_em") or raw.get("created_at"),
        "atualizado_em": raw.get("atualizado_em") or raw.get("updated_at"),
        "deleted_at": raw.get("deleted_at"),
        "persistencia": raw.get("persistencia", "supabase"),
    }


def _normalizar_area(area: dict[str, Any]) -> dict[str, Any]:
    tipo_geometria, geometria_ativa = _geometria_preferencial(area)
    resumo_esboco = area.get("resumo_esboco") or _resumo_vertices(area.get("geometria_esboco"))
    resumo_final = area.get("resumo_final") or _resumo_vertices(area.get("geometria_final"))
    resumo_ativo = resumo_final if tipo_geometria == "final" else resumo_esboco
    status_geometria = _status_geometria(area)
    participantes = area.get("participantes_area") or area.get("participantes") or []

    return {
        **area,
        "status_geometria": status_geometria,
        "tipo_geometria_ativa": tipo_geometria if geometria_ativa else None,
        "geometria_ativa": geometria_ativa,
        "resumo_esboco": resumo_esboco,
        "resumo_final": resumo_final,
        "resumo_ativo": resumo_ativo,
        "anexos": area.get("anexos") or [],
        "codigo_lote": area.get("codigo_lote"),
        "quadra": area.get("quadra"),
        "setor": area.get("setor"),
        "identificacao_lote": _identificacao_lote(area),
        "participantes_area": participantes,
        "participantes_total": len(participantes),
        "status_operacional": _normalizar_status_operacional(area.get("status_operacional"), possui_participante=bool(participantes or area.get("cliente_id")), status_geometria=status_geometria),
        "status_documental": _normalizar_status_documental(area.get("status_documental"), possui_participante=bool(participantes or area.get("cliente_id"))),
    }


def _enriquecer_areas_com_participantes(areas: list[dict[str, Any]], sb=None, participantes_projeto: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    if not areas:
        return []
    cliente = sb or _get_supabase()
    try:
        mapa = listar_participantes_area(cliente, areas, participantes_projeto=participantes_projeto)
    except Exception as exc:
        logger.warning("Falha ao listar participantes por area: %s", exc)
        mapa = {}

    enriquecidas: list[dict[str, Any]] = []
    for area in areas:
        participantes = mapa.get(str(area.get("id")), [])
        enriquecidas.append(_normalizar_area({**area, "participantes_area": participantes, "participantes": participantes}))
    return enriquecidas


def _substituir_placeholders_texto(texto: str, campos: dict[str, Any]) -> str:
    resultado = texto
    for chave, valor in campos.items():
        resultado = resultado.replace(f"{{{{ {chave} }}}}", str(valor))
        resultado = resultado.replace(f"{{{{{chave}}}}}", str(valor))
    return resultado


def _substituir_placeholders_docx(documento, campos: dict[str, Any]) -> None:
    def atualizar_paragrafo(paragrafo) -> None:
        texto = _substituir_placeholders_texto(paragrafo.text, campos)
        if texto == paragrafo.text:
            return
        if paragrafo.runs:
            paragrafo.runs[0].text = texto
            for run in paragrafo.runs[1:]:
                run.text = ""
        else:
            paragrafo.add_run(texto)

    for paragrafo in documento.paragraphs:
        atualizar_paragrafo(paragrafo)

    for tabela in documento.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    atualizar_paragrafo(paragrafo)


def _gerar_carta_confrontacao_docx(
    *,
    projeto: dict[str, Any],
    area_a: dict[str, Any],
    area_b: dict[str, Any],
    confronto: dict[str, Any],
) -> bytes | None:
    if Document is None or not TEMPLATE_CARTA_CONFRONTACAO.exists():
        return None

    documento = Document(str(TEMPLATE_CARTA_CONFRONTACAO))
    campos = {
        "projeto_nome": projeto.get("projeto_nome") or projeto.get("nome") or "Projeto",
        "municipio": projeto.get("municipio") or area_a.get("municipio") or area_b.get("municipio") or "",
        "comarca": projeto.get("comarca") or area_a.get("comarca") or area_b.get("comarca") or "",
        "data_extenso": _data_extenso_ptbr(),
        "proprietario_a": area_a.get("proprietario_nome") or "Nao informado",
        "cpf_a": area_a.get("cpf") or area_a.get("proprietario_cpf") or "",
        "imovel_a": area_a.get("nome") or "Area sem nome",
        "matricula_a": area_a.get("matricula") or "",
        "proprietario_b": area_b.get("proprietario_nome") or "Nao informado",
        "cpf_b": area_b.get("cpf") or area_b.get("proprietario_cpf") or "",
        "imovel_b": area_b.get("nome") or "Area sem nome",
        "matricula_b": area_b.get("matricula") or "",
        "tipo_relacao": confronto.get("tipo") or "divisa",
        "comprimento_contato": f"{confronto.get('contato_m', 0)} m",
        "area_sobreposicao": f"{confronto.get('area_intersecao_ha', 0)} ha",
    }
    _substituir_placeholders_docx(documento, campos)

    buffer = io.BytesIO()
    documento.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def listar_areas_projeto(projeto_id: str, sb=None, participantes_projeto: list[dict[str, Any]] | None = None) -> list[dict[str, Any]]:
    cliente = sb or _get_supabase()
    try:
        res = (
            cliente.table("areas_projeto")
            .select("*")
            .eq("projeto_id", projeto_id)
            .is_("deleted_at", "null")
            .order("atualizado_em", desc=True)
            .execute()
        )
        areas = [_row_para_area(area) or {} for area in (res.data or [])]
        return _enriquecer_areas_com_participantes(areas, sb=cliente, participantes_projeto=participantes_projeto)
    except Exception as exc:
        logger.warning("Falha ao listar areas_projeto no Supabase: %s", exc)
        _falhar_supabase(exc, "listar areas_projeto")


def obter_area(area_id: str, sb=None, participantes_projeto: list[dict[str, Any]] | None = None) -> dict[str, Any] | None:
    cliente = sb or _get_supabase()
    try:
        res = (
            cliente.table("areas_projeto")
            .select("*")
            .eq("id", area_id)
            .is_("deleted_at", "null")
            .maybe_single()
            .execute()
        )
        area = _row_para_area(res.data)
        if not area:
            return None
        return _enriquecer_areas_com_participantes([area], sb=cliente, participantes_projeto=participantes_projeto)[0]
    except Exception as exc:
        logger.warning("Falha ao obter area no Supabase: %s", exc)
        _falhar_supabase(exc, "obter area")


def salvar_area_projeto(
    *,
    projeto_id: str,
    cliente_id: str | None,
    nome: str,
    proprietario_nome: str | None = None,
    municipio: str | None = None,
    estado: str | None = None,
    comarca: str | None = None,
    matricula: str | None = None,
    ccir: str | None = None,
    car: str | None = None,
    observacoes: str | None = None,
    codigo_lote: str | None = None,
    quadra: str | None = None,
    setor: str | None = None,
    status_operacional: str | None = None,
    status_documental: str | None = None,
    origem_tipo: str = "formulario",
    geometria_esboco: list[dict[str, Any]] | None = None,
    geometria_final: list[dict[str, Any]] | None = None,
    anexos: list[dict[str, Any]] | None = None,
    participantes_area: list[dict[str, Any]] | None = None,
    area_id: str | None = None,
    sb=None,
) -> dict[str, Any]:
    cliente = sb or _get_supabase()
    existente = obter_area(area_id, sb=cliente) if area_id else None

    vertices_esboco = (
        _vertices_validos(geometria_esboco)
        if geometria_esboco is not None
        else _vertices_validos((existente or {}).get("geometria_esboco"))
    )
    vertices_final = (
        _vertices_validos(geometria_final)
        if geometria_final is not None
        else _vertices_validos((existente or {}).get("geometria_final"))
    )

    cliente_participante = next((item.get("cliente_id") for item in (participantes_area or []) if item.get("cliente_id")), None)
    cliente_id_efetivo = cliente_id or cliente_participante or (existente or {}).get("cliente_id")
    possui_participante = bool(participantes_area or (existente or {}).get("participantes_area") or cliente_id_efetivo)
    status_geometria = "geometria_final" if vertices_final else ("apenas_esboco" if vertices_esboco else None)
    payload = {
        "id": area_id or (existente or {}).get("id") or str(uuid4()),
        "projeto_id": projeto_id,
        "cliente_id": cliente_id_efetivo,
        "nome": (nome or "").strip() or "Area sem nome",
        "proprietario_nome": proprietario_nome,
        "municipio": municipio,
        "estado": estado,
        "comarca": comarca,
        "matricula": matricula,
        "ccir": ccir,
        "car": car,
        "observacoes": observacoes,
        "codigo_lote": codigo_lote if codigo_lote is not None else (existente or {}).get("codigo_lote"),
        "quadra": quadra if quadra is not None else (existente or {}).get("quadra"),
        "setor": setor if setor is not None else (existente or {}).get("setor"),
        "status_operacional": _normalizar_status_operacional(status_operacional if status_operacional is not None else (existente or {}).get("status_operacional"), possui_participante=possui_participante, status_geometria=status_geometria),
        "status_documental": _normalizar_status_documental(status_documental if status_documental is not None else (existente or {}).get("status_documental"), possui_participante=possui_participante),
        "origem_tipo": origem_tipo,
        "geometria_esboco": vertices_esboco,
        "geometria_final": vertices_final,
        "resumo_esboco": _resumo_vertices(vertices_esboco),
        "resumo_final": _resumo_vertices(vertices_final),
        "anexos": anexos if anexos is not None else ((existente or {}).get("anexos") or []),
        "deleted_at": None,
    }

    try:
        if existente and existente.get("persistencia") != "arquivo_local":
            res = cliente.table("areas_projeto").update(payload).eq("id", payload["id"]).execute()
        else:
            res = cliente.table("areas_projeto").insert(payload).execute()

        registro = None
        if res.data:
            registro = _row_para_area(res.data[0])
        if not registro:
            registro = _row_para_area(
                cliente.table("areas_projeto").select("*").eq("id", payload["id"]).maybe_single().execute().data
            )
        if not registro:
            raise RuntimeError("Falha ao persistir area do projeto no Supabase.")
        if participantes_area is not None:
            try:
                participantes_salvos = salvar_participantes_area(cliente, registro["id"], participantes_area)
                registro["participantes_area"] = participantes_salvos
            except Exception as exc:
                logger.warning("Falha ao persistir participantes da area: %s", exc)
        return _enriquecer_areas_com_participantes([registro], sb=cliente)[0]
    except Exception as exc:
        logger.warning("Falha ao persistir area do projeto no Supabase: %s", exc)
        _falhar_supabase(exc, "persistir area do projeto")


def anexar_arquivos_area(
    *,
    area_id: str,
    cliente_id: str | None,
    arquivos: list[tuple[str, bytes, str | None]],
    sb=None,
) -> list[dict[str, Any]]:
    if not arquivos:
        return []

    cliente = sb or _get_supabase()
    area = obter_area(area_id, sb=cliente)
    if not area:
        return []

    pasta_cliente = UPLOADS_DIR / (cliente_id or "sem-cliente") / area_id
    pasta_cliente.mkdir(parents=True, exist_ok=True)
    anexos = list(area.get("anexos") or [])
    total_arquivos = 0
    total_bytes = 0

    for nome_original, conteudo, content_type in arquivos:
        validar_upload_formulario(nome_original, conteudo)
        total_arquivos += 1
        total_bytes += len(conteudo)
        validar_lote_uploads(total_arquivos=total_arquivos, total_bytes=total_bytes)
        extensao = Path(nome_original or "arquivo.bin").suffix or ".bin"
        nome_seguro = f"{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}_{uuid4().hex[:8]}{extensao}"
        destino = pasta_cliente / nome_seguro
        destino.write_bytes(conteudo)
        anexos.append(
            {
                "id": str(uuid4()),
                "nome_original": nome_original,
                "arquivo_nome": nome_seguro,
                "content_type": content_type,
                "tamanho_bytes": len(conteudo),
                "caminho_local": str(destino),
                "criado_em": _agora_iso(),
            }
        )

    try:
        cliente.table("areas_projeto").update({"anexos": anexos}).eq("id", area_id).execute()
    except Exception as exc:
        logger.warning("Falha ao anexar arquivos da area no Supabase: %s", exc)
        _falhar_supabase(exc, "anexar arquivos da area")
    return anexos


def sintetizar_areas_do_projeto(
    *,
    projeto: dict[str, Any],
    cliente: dict[str, Any] | None,
    perimetro_ativo: dict[str, Any] | None,
    geometria_referencia: dict[str, Any] | None,
    sb=None,
    participantes_projeto: list[dict[str, Any]] | None = None,
) -> list[dict[str, Any]]:
    areas = listar_areas_projeto(projeto.get("id"), sb=sb, participantes_projeto=participantes_projeto)
    if areas:
        return areas

    sinteticas: list[dict[str, Any]] = []
    projeto_id = projeto.get("id")
    cliente_id = cliente.get("id") if cliente else projeto.get("cliente_id")
    proprietario = (cliente or {}).get("nome") or projeto.get("cliente_nome")
    nome_base = projeto.get("nome_imovel") or projeto.get("projeto_nome") or "Area principal"

    if geometria_referencia and _vertices_validos(geometria_referencia.get("vertices")):
        sinteticas.append(
            _normalizar_area(
                {
                    "id": f"{projeto_id}-ref",
                    "projeto_id": projeto_id,
                    "cliente_id": cliente_id,
                    "nome": f"{nome_base} - Esboco",
                    "proprietario_nome": proprietario,
                    "municipio": projeto.get("municipio"),
                    "estado": projeto.get("estado"),
                    "comarca": projeto.get("comarca"),
                    "matricula": projeto.get("matricula"),
                    "status_operacional": "croqui_recebido",
                    "status_documental": "formulario_ok" if cliente_id else "pendente",
                    "origem_tipo": geometria_referencia.get("origem_tipo") or "referencia_cliente",
                    "geometria_esboco": geometria_referencia.get("vertices") or [],
                    "geometria_final": [],
                    "anexos": geometria_referencia.get("anexos") or [],
                    "criado_em": geometria_referencia.get("atualizado_em"),
                    "atualizado_em": geometria_referencia.get("atualizado_em"),
                }
            )
        )

    if perimetro_ativo and _vertices_validos(perimetro_ativo.get("vertices")):
        sinteticas.append(
            _normalizar_area(
                {
                    "id": f"{projeto_id}-tec",
                    "projeto_id": projeto_id,
                    "cliente_id": cliente_id,
                    "nome": f"{nome_base} - Tecnico",
                    "proprietario_nome": proprietario,
                    "municipio": projeto.get("municipio"),
                    "estado": projeto.get("estado"),
                    "comarca": projeto.get("comarca"),
                    "matricula": projeto.get("matricula"),
                    "origem_tipo": "perimetro_tecnico",
                    "geometria_esboco": [],
                    "geometria_final": perimetro_ativo.get("vertices") or [],
                    "anexos": [],
                    "criado_em": perimetro_ativo.get("criado_em"),
                    "atualizado_em": perimetro_ativo.get("criado_em"),
                }
            )
        )

    return _enriquecer_areas_com_participantes(sinteticas, sb=sb, participantes_projeto=participantes_projeto)


def detectar_confrontacoes(areas: list[dict[str, Any]]) -> list[dict[str, Any]]:
    confrontacoes: list[dict[str, Any]] = []
    areas_validas: list[dict[str, Any]] = []
    for area in areas:
        _, vertices = _geometria_preferencial(area)
        polygon = _polygon_from_vertices(vertices)
        if polygon is None:
            continue
        areas_validas.append({**area, "_polygon": polygon})

    for indice, area_a in enumerate(areas_validas):
        for area_b in areas_validas[indice + 1:]:
            poly_a = area_a["_polygon"]
            poly_b = area_b["_polygon"]
            if not poly_a.intersects(poly_b) and not poly_a.touches(poly_b):
                continue

            poly_a_m = _transformar_para_metros(poly_a)
            poly_b_m = _transformar_para_metros(poly_b)
            inter = poly_a_m.intersection(poly_b_m)
            contato = poly_a_m.boundary.intersection(poly_b_m.boundary)

            area_intersecao = float(inter.area) if not inter.is_empty else 0.0
            contato_m = float(contato.length) if not contato.is_empty else 0.0
            tipo = "sobreposicao" if area_intersecao > 0.01 else "divisa"

            confrontacoes.append(
                {
                    "id": f"{area_a['id']}::{area_b['id']}",
                    "tipo": tipo,
                    "status": "detectada",
                    "status_revisao": "detectada",
                    "tipo_relacao": "interna",
                    "origem": "geometria",
                    "revisao_pendente": True,
                    "area_a": {
                        "id": area_a.get("id"),
                        "nome": area_a.get("nome"),
                        "proprietario_nome": area_a.get("proprietario_nome"),
                    },
                    "area_b": {
                        "id": area_b.get("id"),
                        "nome": area_b.get("nome"),
                        "proprietario_nome": area_b.get("proprietario_nome"),
                    },
                    "contato_m": round(contato_m, 2),
                    "area_intersecao_ha": round(area_intersecao / 10000, 4),
                }
            )

    confrontacoes.sort(key=lambda item: (item["tipo"] != "sobreposicao", item["area_a"]["nome"] or ""))
    return confrontacoes


def listar_revisoes_confrontacao(sb, projeto_id: str) -> dict[str, dict[str, Any]]:
    try:
        resposta = (
            sb.table("confrontacoes_revisadas")
            .select("*")
            .eq("projeto_id", projeto_id)
            .is_("deleted_at", "null")
            .execute()
        )
        itens = getattr(resposta, "data", None) or []
    except Exception as exc:
        if "confrontacoes_revisadas" in str(exc).lower():
            return {}
        raise
    return {str(item.get("confronto_id")): item for item in itens if item.get("confronto_id")}


def aplicar_revisoes_confrontacao(confrontacoes: list[dict[str, Any]], revisoes: dict[str, dict[str, Any]]) -> list[dict[str, Any]]:
    enriquecidas: list[dict[str, Any]] = []
    for confronto in confrontacoes:
        revisao = revisoes.get(str(confronto.get("id"))) or {}
        status_revisao = revisao.get("status_revisao") or confronto.get("status_revisao") or confronto.get("status") or "detectada"
        tipo_relacao = revisao.get("tipo_relacao") or confronto.get("tipo_relacao") or "interna"
        enriquecidas.append({
            **confronto,
            "status": status_revisao,
            "status_revisao": status_revisao,
            "tipo_relacao": tipo_relacao,
            "observacao": revisao.get("observacao"),
            "autor_revisao": revisao.get("autor"),
            "revisado_em": revisao.get("atualizado_em") or revisao.get("criado_em"),
            "revisao_pendente": status_revisao == "detectada",
        })
    return enriquecidas


def salvar_revisoes_confrontacao(sb, projeto_id: str, revisoes: list[dict[str, Any]]) -> list[dict[str, Any]]:
    salvas: list[dict[str, Any]] = []
    existentes = listar_revisoes_confrontacao(sb, projeto_id)
    for revisao in revisoes:
        confronto_id = str(revisao.get("confronto_id") or "").strip()
        if not confronto_id:
            continue
        payload = {
            "projeto_id": projeto_id,
            "confronto_id": confronto_id,
            "tipo_relacao": revisao.get("tipo_relacao") or "interna",
            "status_revisao": revisao.get("status_revisao") or "detectada",
            "observacao": revisao.get("observacao"),
            "autor": revisao.get("autor"),
            "deleted_at": None,
        }
        existente = existentes.get(confronto_id)
        try:
            if existente and existente.get("id"):
                resposta = (
                    sb.table("confrontacoes_revisadas")
                    .update(payload)
                    .eq("id", existente.get("id"))
                    .execute()
                )
            else:
                resposta = sb.table("confrontacoes_revisadas").insert(payload).execute()
            dados = getattr(resposta, "data", None) or []
            salvas.append(dados[0] if dados else {**payload, **({"id": (existente or {}).get("id")} if existente else {})})
        except Exception as exc:
            if "confrontacoes_revisadas" in str(exc).lower():
                salvas.append(payload)
                continue
            raise
    return salvas


def gerar_cartas_confrontacao_zip(
    *,
    projeto: dict[str, Any],
    areas: list[dict[str, Any]],
    confrontacoes: list[dict[str, Any]],
) -> bytes:
    mapa_areas = {area.get("id"): area for area in areas}
    zip_buffer = io.BytesIO()
    nome_projeto = projeto.get("projeto_nome") or projeto.get("nome") or "Projeto"

    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        if not confrontacoes:
            zf.writestr(
                "LEIA-ME.txt",
                (
                    f"Nenhuma confrontacao geometrica foi detectada para {nome_projeto}.\n"
                    "Cadastre mais areas ou confirme confrontantes manualmente no app.\n"
                ).encode("utf-8"),
            )
        for indice, confronto in enumerate(confrontacoes, start=1):
            area_a = mapa_areas.get(confronto["area_a"]["id"], confronto["area_a"])
            area_b = mapa_areas.get(confronto["area_b"]["id"], confronto["area_b"])
            docx_bytes = _gerar_carta_confrontacao_docx(
                projeto=projeto,
                area_a=area_a,
                area_b=area_b,
                confronto=confronto,
            )
            if docx_bytes:
                nome_docx = f"CARTA_CONFRONTACAO_{indice:02d}.docx"
                zf.writestr(nome_docx, docx_bytes)
                continue

            texto = f"""
CARTA DE CONFRONTACAO
=====================
Projeto: {nome_projeto}

Area principal:
- Nome: {area_a.get('nome') or 'Area sem nome'}
- Proprietario: {area_a.get('proprietario_nome') or 'Nao informado'}
- Matricula: {area_a.get('matricula') or 'Nao informada'}

Area confrontante:
- Nome: {area_b.get('nome') or 'Area sem nome'}
- Proprietario: {area_b.get('proprietario_nome') or 'Nao informado'}
- Matricula: {area_b.get('matricula') or 'Nao informada'}

Resumo tecnico:
- Tipo de relacao: {confronto.get('tipo')}
- Comprimento aproximado de contato: {confronto.get('contato_m', 0)} m
- Sobreposicao aproximada: {confronto.get('area_intersecao_ha', 0)} ha

Observacao:
Esta carta foi preparada automaticamente a partir das geometrias disponiveis
no GeoAdmin. Revise os dados pessoais, o esboco do cliente e o perimetro
tecnico antes da assinatura definitiva.
""".strip()
            nome_txt = f"CARTA_CONFRONTACAO_{indice:02d}.txt"
            zf.writestr(nome_txt, texto.encode("utf-8"))

        zf.writestr(
            "manifesto_cartas.json",
            json.dumps(
                {
                    "gerado_em": _agora_iso(),
                    "projeto": {
                        "id": projeto.get("id"),
                        "nome": nome_projeto,
                    },
                    "areas_total": len(areas),
                    "confrontacoes_total": len(confrontacoes),
                },
                ensure_ascii=False,
                indent=2,
            ).encode("utf-8"),
        )

    zip_buffer.seek(0)
    return zip_buffer.getvalue()
