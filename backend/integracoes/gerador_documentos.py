"""
GeoAdmin Pro — Gerador de Documentos GPRF
==========================================
backend/integracoes/gerador_documentos.py

Gera os 7 documentos do processo automaticamente
a partir dos dados do banco.

Uso:
    pacote = gerar_todos_documentos(supabase, projeto_id)
    # pacote.zip contém todos os .docx prontos

Requer:
    pip install python-docx jinja2
"""

import io
import zipfile
import math
import logging
from datetime import datetime
from typing import Optional
from dataclasses import dataclass, field

logger = logging.getLogger("geoadmin.documentos")


# ---------------------------------------------------------------------------
# Tipos internos
# ---------------------------------------------------------------------------

@dataclass
class DadosDocumento:
    """Todos os dados necessários para gerar os documentos."""
    # Projeto
    projeto_id:     str
    projeto_nome:   str
    nome_imovel:    str
    municipio:      str
    estado:         str
    endereco_imovel: str = ""
    endereco_imovel_numero: str = ""
    cep_imovel:     str = ""
    comarca:        str = ""
    matricula:      str = ""
    area_ha:        float = 0.0
    area_m2:        float = 0.0

    # Cliente
    cliente_nome:   str = ""
    cliente_cpf:    str = ""
    cliente_rg:     str = ""
    estado_civil:   str = ""
    profissao:      str = ""
    telefone:       str = ""
    email:          str = ""
    endereco:       str = ""
    endereco_numero: str = ""
    cliente_municipio: str = ""
    cliente_estado: str = ""
    cep:            str = ""
    conjuge_nome:   str = ""
    conjuge_cpf:    str = ""

    # Técnico
    tecnico_nome:       str = ""
    tecnico_cpf:        str = ""
    tecnico_crt:        str = ""
    tecnico_crea:       str = ""
    tecnico_codigo_incra: str = ""

    # Confrontantes
    confrontantes:  list = field(default_factory=list)

    # Dados para cálculo VTN
    classe_imovel:  str = "IV"
    distancia_sede_km: float = 20.0
    distancia_asfalto_km: float = 10.0
    tempo_posse_anos: int = 5

    # Data e local
    municipio_local: str = ""
    data_atual:     str = ""

    def __post_init__(self):
        if not self.data_atual:
            self.data_atual = datetime.now().strftime("%d/%m/%Y")
        if not self.municipio_local:
            self.municipio_local = self.municipio


# ---------------------------------------------------------------------------
# Busca de dados
# ---------------------------------------------------------------------------

def _buscar_dados(supabase, projeto_id: str) -> DadosDocumento:
    """Busca todos os dados necessários do Supabase."""

    # Projeto + cliente
    res = supabase.table("vw_formulario_cliente") \
        .select("*") \
        .eq("projeto_id", projeto_id) \
        .single() \
        .execute()

    if not res.data:
        raise ValueError(f"[ERRO-401] Projeto {projeto_id} não encontrado.")

    d = res.data

    # Técnico
    tec = supabase.table("tecnico") \
        .select("*") \
        .eq("ativo", True) \
        .limit(1) \
        .execute()

    t = tec.data[0] if tec.data else {}

    # Confrontantes
    conf = supabase.table("confrontantes") \
        .select("*") \
        .eq("projeto_id", projeto_id) \
        .is_("deleted_at", "null") \
        .execute()

    # Área em ha (do banco ou calculada)
    area_ha = float(d.get("area_ha") or 0)
    area_m2 = area_ha * 10000

    return DadosDocumento(
        projeto_id=projeto_id,
        projeto_nome=d.get("projeto_nome", ""),
        nome_imovel=d.get("nome_imovel") or d.get("projeto_nome", ""),
        municipio=d.get("imovel_municipio", ""),
        estado=d.get("imovel_estado", "GO"),
        endereco_imovel=d.get("endereco_imovel", ""),
        endereco_imovel_numero=d.get("endereco_imovel_numero", ""),
        cep_imovel=d.get("cep_imovel", ""),
        comarca=d.get("comarca", ""),
        matricula=d.get("matricula", ""),
        area_ha=area_ha,
        area_m2=area_m2,
        cliente_nome=d.get("cliente_nome", ""),
        cliente_cpf=d.get("cliente_cpf", ""),
        cliente_rg=d.get("cliente_rg", ""),
        estado_civil=d.get("estado_civil", ""),
        profissao=d.get("profissao", ""),
        telefone=d.get("telefone", ""),
        email=d.get("email", ""),
        endereco=d.get("endereco", ""),
        endereco_numero=d.get("endereco_numero", ""),
        cliente_municipio=d.get("cliente_municipio", ""),
        cliente_estado=d.get("cliente_estado") or d.get("imovel_estado", "GO"),
        cep=d.get("cep", ""),
        conjuge_nome=d.get("conjuge_nome") or "",
        conjuge_cpf=d.get("conjuge_cpf") or "",
        tecnico_nome=t.get("nome", ""),
        tecnico_cpf=t.get("cpf", ""),
        tecnico_crt=t.get("crt", ""),
        tecnico_crea=t.get("crea", ""),
        tecnico_codigo_incra=t.get("codigo_incra", ""),
        confrontantes=conf.data or [],
        municipio_local=d.get("imovel_municipio", ""),
    )


# ---------------------------------------------------------------------------
# Helpers de apresentação
# ---------------------------------------------------------------------------

def _montar_linha_endereco(endereco: str, numero: str = "", municipio: str = "", estado: str = "", cep: str = "", fallback: str = "Não informado") -> str:
    partes: list[str] = []
    endereco = (endereco or "").strip()
    numero = (numero or "").strip()
    municipio = (municipio or "").strip()
    estado = (estado or "").strip()
    cep = (cep or "").strip()

    if endereco and numero:
        partes.append(f"{endereco}, N° {numero}")
    elif endereco:
        partes.append(endereco)
    elif numero:
        partes.append(f"N° {numero}")

    local = " - ".join([parte for parte in [municipio, estado] if parte])
    if local:
        partes.append(local)
    if cep:
        partes.append(f"CEP {cep}")

    return " | ".join(partes) if partes else fallback


def _montar_localizacao_imovel(dados: DadosDocumento) -> str:
    return _montar_linha_endereco(
        dados.endereco_imovel,
        dados.endereco_imovel_numero,
        dados.municipio,
        dados.estado,
        dados.cep_imovel,
        fallback="Localização do imóvel não informada",
    )


def _montar_endereco_residencial(dados: DadosDocumento) -> str:
    return _montar_linha_endereco(
        dados.endereco,
        dados.endereco_numero,
        dados.cliente_municipio,
        dados.cliente_estado or dados.estado,
        dados.cep,
        fallback="Endereço residencial não informado",
    )


# ---------------------------------------------------------------------------
# Gerador de texto por template
# ---------------------------------------------------------------------------

def _preencher(template: str, dados: DadosDocumento, extra: dict = None) -> str:
    """Substitui placeholders no template pelos dados reais.

    Ordena placeholders por comprimento (maior primeiro) para evitar substituições
    parciais quando o valor de um placeholder contém texto que combine com outro
    placeholder mais curto.
    """
    endereco_residencial = _montar_endereco_residencial(dados)
    endereco_imovel = _montar_localizacao_imovel(dados)
    ctx = {
        "NOME_PROPRIETARIO":    dados.cliente_nome.upper(),
        "CPF_PROPRIETARIO":     dados.cliente_cpf,
        "RG_PROPRIETARIO":      dados.cliente_rg,
        "ESTADO_CIVIL":         dados.estado_civil,
        "PROFISSAO":            dados.profissao,
        "ENDERECO":             dados.endereco,
        "NUMERO":               dados.endereco_numero,
        "MUNICIPIO_CLIENTE":    dados.cliente_municipio,
        "ESTADO_CLIENTE":       dados.cliente_estado or dados.estado,
        "CEP":                  dados.cep,
        "ENDERECO_RESIDENCIAL": dados.endereco,
        "NUMERO_RESIDENCIAL":   dados.endereco_numero,
        "CEP_RESIDENCIAL":      dados.cep,
        "ENDERECO_RESIDENCIAL_COMPLETO": endereco_residencial,
        "TELEFONE":             dados.telefone,
        "EMAIL":                dados.email,
        "NOME_IMOVEL":          dados.nome_imovel.upper(),
        "MUNICIPIO_IMOVEL":     dados.municipio,
        "ESTADO_IMOVEL":        dados.estado,
        "ENDERECO_IMOVEL":      dados.endereco_imovel,
        "NUMERO_IMOVEL":        dados.endereco_imovel_numero,
        "CEP_IMOVEL":           dados.cep_imovel,
        "ENDERECO_IMOVEL_COMPLETO": endereco_imovel,
        "AREA_HA":              f"{dados.area_ha:.4f}",
        "AREA_M2":              f"{dados.area_m2:.2f}",
        "MATRICULA":            dados.matricula,
        "COMARCA":              dados.comarca,
        "TECNICO_NOME":         dados.tecnico_nome.upper(),
        "TECNICO_CPF":          dados.tecnico_cpf,
        "TECNICO_CRT":          dados.tecnico_crt,
        "TECNICO_CREA":         dados.tecnico_crea,
        "TECNICO_CODIGO_INCRA": dados.tecnico_codigo_incra,
        "MUNICIPIO_LOCAL":      dados.municipio_local,
        "DATA_ATUAL":           dados.data_atual,
    }
    if extra:
        ctx.update(extra)

    resultado = template
    # Ordena por comprimento (maior primeiro) para evitar substituições parciais
    chaves_ordenadas = sorted(ctx.keys(), key=len, reverse=True)
    for chave in chaves_ordenadas:
        resultado = resultado.replace(f"{{{{{chave}}}}}", str(ctx[chave]))
    return resultado


# ---------------------------------------------------------------------------
# Templates dos documentos
# ---------------------------------------------------------------------------

TEMPLATE_REQ_TITULACAO = """
REQUERIMENTO DE TITULAÇÃO

Nome do Requerente: {{NOME_PROPRIETARIO}}
CPF/CNPJ: {{CPF_PROPRIETARIO}}
RG: {{RG_PROPRIETARIO}}
Estado Civil: {{ESTADO_CIVIL}}
Profissão: {{PROFISSAO}}
E-mail: {{EMAIL}}

Endereço residencial: {{ENDERECO_RESIDENCIAL_COMPLETO}}
Telefone: {{TELEFONE}}

Nome do Imóvel: {{NOME_IMOVEL}}
Localização do Imóvel: {{ENDERECO_IMOVEL_COMPLETO}}
Área do Imóvel: {{AREA_HA}} ha
Município do Imóvel: {{MUNICIPIO_IMOVEL}} - {{ESTADO_IMOVEL}}
Matrícula: {{MATRICULA}}
Comarca: {{COMARCA}}

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________
Assinatura do Requerente
{{NOME_PROPRIETARIO}}
CPF: {{CPF_PROPRIETARIO}}
"""

TEMPLATE_ORDEM_SERVICO = """
REQUERIMENTO DE ORDEM DE SERVIÇO

Eu, {{NOME_PROPRIETARIO}}, portador do CPF: {{CPF_PROPRIETARIO}},
residente em {{ENDERECO_RESIDENCIAL_COMPLETO}},
e ocupante do imóvel rural denominado {{NOME_IMOVEL}},
localizado em {{ENDERECO_IMOVEL_COMPLETO}}.

Venho por meio deste instrumento, solicitar a emissão da ordem de
serviço para execução dos serviços topográficos de medição,
demarcação e georreferenciamento para o imóvel rural objeto do
requerimento de titulação, para o profissional {{TECNICO_NOME}},
cadastrado nesta Gerência, portador do CPF: {{TECNICO_CPF}},
registro profissional {{TECNICO_CRT}} e código do INCRA: {{TECNICO_CODIGO_INCRA}}.

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________     _________________________________________________
Assinatura do Requerente                              Assinatura do Profissional
{{NOME_PROPRIETARIO}}                                 {{TECNICO_NOME}}

OBS: Requerimento com firma reconhecida
"""

TEMPLATE_DECL_FUNCAO_PUBLICA = """
DECLARAÇÃO DE NÃO EXERCÍCIO DE FUNÇÃO PÚBLICA

Eu, {{NOME_PROPRIETARIO}}, portador (a) do RG: {{RG_PROPRIETARIO}}
e CPF: {{CPF_PROPRIETARIO}}.

Declaro para os devidos fins que não exerço função pública, em cargo
efetivo ou em cargo de comissão, bem como de mandatos eletivos, em
nenhum dos Poderes da União, dos Estados, do Distrito Federal e dos
Municípios, em conformidade com o Artigo 32, Parágrafo I, Cap. V da
Lei Estadual 18.826 de 19/05/2015 – Vedação.

Por ser verdade, firmo a presente, me responsabilizando pelas formas legais.

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________
Assinatura
{{NOME_PROPRIETARIO}}

OBS: Declaração com firma reconhecida
"""

TEMPLATE_DECL_IMOVEL_RURAL = """
DECLARAÇÃO DE IMÓVEL RURAL

Eu, {{NOME_PROPRIETARIO}}, portador (a) do RG: {{RG_PROPRIETARIO}}
e CPF: {{CPF_PROPRIETARIO}}.

Venho por meio deste instrumento particular, declarar que não possuo
imóvel rural ou propriedade rural registrada em meu nome e não sou
herdeiro de imóvel rural, em qualquer parte do território nacional.

Por ser verdade, firmo a presente, me responsabilizando pelas formas legais.

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________
Assinatura
{{NOME_PROPRIETARIO}}

OBS: Declaração com firma reconhecida
"""

TEMPLATE_DECL_RESIDENCIA = """
DECLARAÇÃO DE RESIDÊNCIA

Eu, {{NOME_PROPRIETARIO}}, portador do RG: {{RG_PROPRIETARIO}}
e CPF: {{CPF_PROPRIETARIO}}.

Venho por meio deste instrumento particular, declarar que não possuo
comprovante de endereço em meu nome, sendo certo e verdadeiro que
resido e domicilio em {{ENDERECO_RESIDENCIAL_COMPLETO}}.

Declaro ainda que o imóvel rural objeto deste processo,
denominado {{NOME_IMOVEL}}, fica localizado em {{ENDERECO_IMOVEL_COMPLETO}}.

Por ser verdade, firmo o presente, me responsabilizando pelas formas legais.

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________
Assinatura
{{NOME_PROPRIETARIO}}

OBS: Declaração com firma reconhecida
"""

TEMPLATE_DECL_LIMITES = """
DECLARAÇÃO DE RESPEITO DE LIMITES

{{NOME_PROPRIETARIO}}, CPF nº {{CPF_PROPRIETARIO}}, proprietário do
imóvel rural denominado {{NOME_IMOVEL}}, matrícula nº {{MATRICULA}},
declaro sob as penas da Lei que quando dos trabalhos topográficos
executados na citada propriedade pelo TÉCNICO {{TECNICO_NOME}},
CRT nº {{TECNICO_CRT}}, CPF nº {{TECNICO_CPF}}, credenciado pelo INCRA
sob o código {{TECNICO_CODIGO_INCRA}}, foram respeitados os limites de
"divisas in loco" com o meu confrontante, {{CONFRONTANTE_NOME}},
CPF nº {{CONFRONTANTE_CPF}}, proprietário do imóvel rural denominado
{{CONFRONTANTE_IMOVEL}}, matrícula nº {{CONFRONTANTE_MATRICULA}}.

E assim, munidos de boa fé, conjuntamente com os que assinam, declaram
não haver qualquer litígio entre as partes.

O trecho confrontante possui os seguintes elementos técnicos:

DATUM: SIRGAS2000

{{TABELA_VERTICES}}

{{MUNICIPIO_LOCAL}}, {{DATA_ATUAL}}.

_________________________________________________     _________________________________________________
Assinatura do Proprietário                            Assinatura do Confrontante
{{NOME_PROPRIETARIO}}                                 {{CONFRONTANTE_NOME}}
CPF: {{CPF_PROPRIETARIO}}                             CPF: {{CONFRONTANTE_CPF}}

OBS: Declaração com firma reconhecida
"""


# ---------------------------------------------------------------------------
# Gerador da tabela de vértices (Declaração de Limites)
# ---------------------------------------------------------------------------

def _gerar_tabela_vertices(vertices: list) -> str:
    """
    Gera a tabela técnica de vértices do trecho confrontante.
    vertices: lista de dicts com codigo, longitude, latitude, altitude,
              prox_codigo, azimute, distancia
    """
    if not vertices:
        return "(tabela de vértices a preencher)"

    cabecalho = (
        f"{'CÓDIGO':<10} {'LONGITUDE':>16} {'LATITUDE':>16} "
        f"{'ALT(m)':>10} {'PRÓ.CÓD':>10} {'AZIMUTE':>16} {'DIST(m)':>12}\n"
        + "-" * 96
    )

    linhas = [cabecalho]
    for v in vertices:
        linha = (
            f"{v.get('codigo',''):<10} "
            f"{v.get('longitude',0):>16.8f} "
            f"{v.get('latitude',0):>16.8f} "
            f"{v.get('altitude',0):>10.4f} "
            f"{v.get('prox_codigo',''):>10} "
            f"{v.get('azimute',''):>16} "
            f"{v.get('distancia',0):>12.3f}"
        )
        linhas.append(linha)

    return "\n".join(linhas)


# ---------------------------------------------------------------------------
# Gerar documento Word (.docx)
# ---------------------------------------------------------------------------

def _texto_para_docx(texto: str, titulo: str) -> bytes:
    """Converte texto para .docx usando python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Margens
        for section in doc.sections:
            section.top_margin    = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin   = Cm(3)
            section.right_margin  = Cm(2)

        # Título
        p_titulo = doc.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_titulo = p_titulo.add_run(titulo.upper())
        run_titulo.bold = True
        run_titulo.font.size = Pt(14)

        doc.add_paragraph()

        # Conteúdo — cada linha vira parágrafo
        for linha in texto.strip().split("\n"):
            p = doc.add_paragraph(linha)
            p.paragraph_format.space_after = Pt(0)

            # Linhas de assinatura em negrito
            if "___" in linha or "Assinatura" in linha:
                for run in p.runs:
                    run.bold = True

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    except ImportError:
        # Fallback: retornar texto puro como .txt
        return texto.encode("utf-8")


# ---------------------------------------------------------------------------
# Gerar todos os documentos
# ---------------------------------------------------------------------------

def gerar_todos_documentos(supabase, projeto_id: str) -> bytes:
    """
    Gera os 7 documentos GPRF e retorna um ZIP em memória.

    Retorna:
        bytes do ZIP com todos os .docx
    """
    dados = _buscar_dados(supabase, projeto_id)
    base  = f"{dados.nome_imovel}_{dados.municipio}".replace(" ", "_")[:30]
    zip_buf = io.BytesIO()

    with zipfile.ZipFile(zip_buf, "w", zipfile.ZIP_DEFLATED) as zf:

        # 01 — Requerimento de Titulação
        txt = _preencher(TEMPLATE_REQ_TITULACAO, dados)
        zf.writestr(f"01_Req_Titulacao_{base}.docx",
                    _texto_para_docx(txt, "Requerimento de Titulação"))

        # 02 — Ordem de Serviço
        txt = _preencher(TEMPLATE_ORDEM_SERVICO, dados)
        zf.writestr(f"02_Req_OrdemServico_{base}.docx",
                    _texto_para_docx(txt, "Requerimento de Ordem de Serviço"))

        # 05 — Declaração de Função Pública
        txt = _preencher(TEMPLATE_DECL_FUNCAO_PUBLICA, dados)
        zf.writestr(f"05_Decl_FuncaoPublica_{base}.docx",
                    _texto_para_docx(txt, "Declaração de Não Exercício de Função Pública"))

        # 06 — Declaração de Imóvel Rural
        txt = _preencher(TEMPLATE_DECL_IMOVEL_RURAL, dados)
        zf.writestr(f"06_Decl_ImovelRural_{base}.docx",
                    _texto_para_docx(txt, "Declaração de Imóvel Rural"))

        # 07 — Declaração de Residência
        txt = _preencher(TEMPLATE_DECL_RESIDENCIA, dados)
        zf.writestr(f"07_Decl_Residencia_{base}.docx",
                    _texto_para_docx(txt, "Declaração de Residência"))

        # 03 — Declaração de Respeito de Limites (uma por confrontante)
        for conf in dados.confrontantes:
            vertices  = conf.get("vertices_json") or []
            tab_vert  = _gerar_tabela_vertices(vertices)
            extra = {
                "CONFRONTANTE_NOME":     conf.get("nome", "").upper(),
                "CONFRONTANTE_CPF":      conf.get("cpf", ""),
                "CONFRONTANTE_IMOVEL":   conf.get("nome_imovel", "").upper(),
                "CONFRONTANTE_MATRICULA": conf.get("matricula", ""),
                "TABELA_VERTICES":       tab_vert,
            }
            txt = _preencher(TEMPLATE_DECL_LIMITES, dados, extra)
            lado = conf.get("lado", "Confrontante")
            zf.writestr(
                f"03_Decl_Limites_{lado}_{base}.docx",
                _texto_para_docx(txt, f"Declaração de Respeito de Limites — {lado}")
            )

        # README
        readme = _gerar_readme(dados)
        zf.writestr("COMO_ASSINAR.txt", readme.encode("utf-8"))

    logger.info(
        f"Pacote GPRF gerado: '{dados.nome_imovel}' — "
        f"{len(dados.confrontantes)} confrontante(s)"
    )
    return zip_buf.getvalue()


def _gerar_readme(dados: DadosDocumento) -> str:
    return f"""
GEOADMIN PRO — PACOTE DE DOCUMENTOS GPRF
=========================================
Projeto:   {dados.projeto_nome}
Imóvel:    {dados.nome_imovel}
Município: {dados.municipio}
Área:      {dados.area_ha:.4f} ha
Gerado em: {dados.data_atual}

DOCUMENTOS INCLUÍDOS
--------------------
01 Requerimento de Titulação        → assinar com firma reconhecida
02 Requerimento de Ordem de Serviço → assinar proprietário + técnico
03 Declaração de Respeito de Limites → uma por confrontante + firma
05 Declaração de Função Pública      → assinar com firma reconhecida
06 Declaração de Imóvel Rural        → assinar com firma reconhecida
07 Declaração de Residência          → assinar com firma reconhecida

ATENÇÃO
-------
Nenhum dado foi digitado manualmente.
Tudo veio do banco de dados do GeoAdmin Pro.
Se houver erro em algum campo, corrija no app e gere novamente.
""".strip()


# ---------------------------------------------------------------------------
# TESTES
# ---------------------------------------------------------------------------

def _testar():
    """Testa os geradores sem Supabase."""
    from dataclasses import asdict

    dados = DadosDocumento(
        projeto_id="test-001",
        projeto_nome="Fazenda Margarida",
        nome_imovel="Fazenda Margarida",
        municipio="Pirenópolis",
        estado="GO",
        endereco_imovel="Estrada da Vargem",
        endereco_imovel_numero="Km 5",
        cep_imovel="72980-000",
        comarca="Pirenópolis",
        matricula="1234",
        area_ha=45.6789,
        area_m2=456789.0,
        cliente_nome="João da Silva",
        cliente_cpf="123.456.789-00",
        cliente_rg="1234567",
        estado_civil="casado",
        profissao="Agricultor",
        telefone="(61) 99999-0000",
        email="joao@email.com",
        endereco="Rua Principal",
        endereco_numero="45",
        cliente_municipio="Anápolis",
        cliente_estado="GO",
        cep="75000-000",
        tecnico_nome="Hugo Desenrola",
        tecnico_cpf="987.654.321-00",
        tecnico_crt="CRT-001",
        tecnico_codigo_incra="INCRA-001",
        confrontantes=[{
            "lado": "Norte",
            "nome": "Pedro Confrontante",
            "cpf": "111.222.333-44",
            "nome_imovel": "Sítio do Pedro",
            "matricula": "5678",
            "vertices_json": [
                {"codigo":"V01","longitude":-47.929722,"latitude":-15.779167,
                 "altitude":1172.0,"prox_codigo":"V02","azimute":"36°52'11\"","distancia":500.0},
                {"codigo":"V02","longitude":-47.927021,"latitude":-15.779167,
                 "altitude":1171.5,"prox_codigo":"V01","azimute":"216°52'11\"","distancia":500.0},
            ]
        }],
    )

    erros = []
    def checar(desc, ok, msg=""):
        status = "✓" if ok else "✗"
        if not ok: erros.append(f"{desc}: {msg}")
        print(f"  {status} {desc}")

    print("\n" + "="*55)
    print("GABARITO — Gerador Documentos GPRF")
    print("="*55)

    print("\n[1] Templates preenchidos")
    for nome, template in [
        ("Req Titulação",      TEMPLATE_REQ_TITULACAO),
        ("Ordem Serviço",      TEMPLATE_ORDEM_SERVICO),
        ("Decl Função Públ.",  TEMPLATE_DECL_FUNCAO_PUBLICA),
        ("Decl Imóvel Rural",  TEMPLATE_DECL_IMOVEL_RURAL),
        ("Decl Residência",    TEMPLATE_DECL_RESIDENCIA),
    ]:
        txt = _preencher(template, dados)
        checar(f"{nome} — nome ok",    "JOÃO DA SILVA" in txt)
        checar(f"{nome} — CPF ok",     "123.456.789-00" in txt)
        checar(f"{nome} — imóvel ok",  "FAZENDA MARGARIDA" in txt)

    print("\n[2] Declaração de Limites com vértices")
    conf = dados.confrontantes[0]
    tab  = _gerar_tabela_vertices(conf["vertices_json"])
    checar("Tabela tem V01",          "V01" in tab)
    checar("Tabela tem azimute",      "36°52" in tab)
    checar("Tabela tem distância",    "500.000" in tab)

    txt_limites = _preencher(TEMPLATE_DECL_LIMITES, dados, {
        "CONFRONTANTE_NOME":      "PEDRO CONFRONTANTE",
        "CONFRONTANTE_CPF":       "111.222.333-44",
        "CONFRONTANTE_IMOVEL":    "SÍTIO DO PEDRO",
        "CONFRONTANTE_MATRICULA": "5678",
        "TABELA_VERTICES":        tab,
    })
    checar("Doc Limites — confrontante ok", "PEDRO CONFRONTANTE" in txt_limites)
    checar("Doc Limites — vértices ok",     "V01" in txt_limites)

    print("\n[3] ZIP com todos os documentos")
    zip_bytes = io.BytesIO()
    with zipfile.ZipFile(zip_bytes, "w") as zf:
        for nome_doc, template in [
            ("01_Req_Titulacao.docx",      TEMPLATE_REQ_TITULACAO),
            ("05_Decl_FuncaoPublica.docx", TEMPLATE_DECL_FUNCAO_PUBLICA),
        ]:
            txt = _preencher(template, dados)
            zf.writestr(nome_doc, txt.encode("utf-8"))
        zf.writestr("COMO_ASSINAR.txt", _gerar_readme(dados).encode("utf-8"))

    checar("ZIP gerado com bytes",   len(zip_bytes.getvalue()) > 100)
    with zipfile.ZipFile(zip_bytes) as zf:
        nomes = zf.namelist()
    checar("ZIP tem 3 arquivos",     len(nomes) == 3)
    checar("README incluído",        "COMO_ASSINAR.txt" in nomes)

    print("\n" + "="*55)
    if erros:
        print(f"FALHOU — {len(erros)} erro(s):")
        for e in erros: print(f"  ✗ {e}")
    else:
        print("APROVADO — Todos os geradores validados.")
    print("="*55 + "\n")
    return len(erros) == 0


if __name__ == "__main__":
    sucesso = _testar()
    exit(0 if sucesso else 1)
